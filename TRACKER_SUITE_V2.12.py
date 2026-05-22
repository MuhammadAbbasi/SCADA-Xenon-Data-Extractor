import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os, csv, glob, re, shutil
import pandas as pd
import threading, multiprocessing
from concurrent.futures import ProcessPoolExecutor, as_completed
import sys, time, gc, queue
from datetime import datetime, timedelta
from collections import Counter

try:
    from tkcalendar import Calendar as TkCalendar
    HAS_TKCALENDAR = True
except ImportError:
    HAS_TKCALENDAR = False

try:
    from astral import LocationInfo
    from astral.sun import sun
    import pytz
    HAS_ASTRAL = True
except ImportError:
    HAS_ASTRAL = False

from PIL import Image, ImageTk
import matplotlib
matplotlib.use("TkAgg")
from matplotlib.figure import Figure
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

# ==============================================================================
# CONSTANTS
# ==============================================================================

ITALIAN_MONTHS = {
    1: "Gennaio", 2: "Febbraio", 3: "Marzo", 4: "Aprile",
    5: "Maggio", 6: "Giugno", 7: "Luglio", 8: "Agosto",
    9: "Settembre", 10: "Ottobre", 11: "Novembre", 12: "Dicembre"
}
DEFAULT_ROOT = r"//S01/get/2025.01 Mazara 01 A2A/03 - REPORT/Report/04 Tracker report"
TEMPLATE_FILENAME = "YYYY.MM.DD - Report Tracker.docx"

# ==============================================================================
# SOLAR CALCULATION
# ==============================================================================

def get_sun_times(date_obj, lat=37.7717, lon=12.6304):
    if not HAS_ASTRAL:
        base = datetime.combine(date_obj.date(), datetime.min.time())
        return base + timedelta(hours=7), base + timedelta(hours=18)
    try:
        site = LocationInfo("Mazara", "Italy", "Europe/Rome", lat, lon)
        s = sun(site.observer, date=date_obj)
        tz = pytz.timezone("Europe/Rome")
        sr = s['sunrise'].astimezone(tz).replace(tzinfo=None)
        ss = s['sunset'].astimezone(tz).replace(tzinfo=None)
        return sr, ss
    except Exception:
        base = datetime.combine(date_obj.date(), datetime.min.time())
        return base + timedelta(hours=7), base + timedelta(hours=18)

# ==============================================================================
# WORKER FUNCTIONS
# ==============================================================================

def worker_extract(args):
    infile, outfile = args
    filename = os.path.basename(infile)
    possible_encodings = ["utf-8-sig", "utf-16", "latin1"]
    lines = None
    for enc in possible_encodings:
        try:
            with open(infile, "r", encoding=enc) as f:
                lines = f.readlines()
            break
        except UnicodeDecodeError:
            continue
    if not lines:
        return f"FAILED:{filename}"
    count = 0
    try:
        with open(outfile, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["Tag", "Name", "Value", "Unit", "Type", "Timestamp"])
            for line in lines:
                if "Angolo" in line or "angolo" in line:
                    parts = [p.strip() for p in line.strip().split(";")]
                    if len(parts) >= 6:
                        name = parts[1].lower()
                        if "angolo target" in name or "angolo attuale" in name:
                            writer.writerow(parts[:6])
                            count += 1
        del lines
        return f"OK:{filename}:{count}"
    except Exception as e:
        return f"ERROR:{filename}:{e}"


def worker_read_csv(filepath):
    filename = os.path.basename(filepath)
    pattern = re.compile(r"Data_Mod_NCU_(\d+)_TCU_(\d+)\.TC(\d+),([^,]+),([^,]+),([^,]+),([^,]+),(.+)")
    rows = []
    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            for line in f:
                m = pattern.match(line.strip())
                if m:
                    rows.append({
                        "NCU": int(m.group(1)), "TCU": int(m.group(2)), "TC": int(m.group(3)),
                        "Parametro": m.group(4), "Valore": float(m.group(5).replace(',', '.')),
                        "Timestamp": m.group(8)
                    })
        return pd.DataFrame(rows)
    except Exception as e:
        print(f"[MERGE] Error reading {filename}: {e}", flush=True)
        return pd.DataFrame()


def worker_plot_file(args):
    import matplotlib.pyplot as plt
    matplotlib.use('Agg')
    ncu, tcu, data_dict, out_path = args
    try:
        df = pd.DataFrame(data_dict)
        df["Timestamp"] = pd.to_datetime(df["Timestamp"])
        target = df[df["TC"] == 4].sort_values("Timestamp")
        actual = df[df["TC"] == 5].sort_values("Timestamp")
        if target.empty and actual.empty:
            return "Empty"
        fig, ax = plt.subplots(figsize=(10, 6))
        if not target.empty:
            ax.plot(target["Timestamp"], target["Valore"], label="Target", color="royalblue", lw=1)
        if not actual.empty:
            ax.plot(actual["Timestamp"], actual["Valore"], label="Actual", color="tomato", lw=1)
        ax.set_title(f"NCU {ncu} - TCU {tcu}")
        ax.set_ylabel("Angle (deg)")
        ax.grid(True, alpha=0.3)
        ax.legend(loc="upper right")
        if not actual.empty:
            stats = f"Min: {actual['Valore'].min():.1f}  Max: {actual['Valore'].max():.1f}  Avg: {actual['Valore'].mean():.1f}"
            ax.text(0.02, 0.02, stats, transform=ax.transAxes, fontsize=9,
                    bbox=dict(boxstyle="round,pad=0.3", facecolor='white', alpha=0.8, edgecolor='#cccccc'))
        fig.tight_layout()
        fig.savefig(out_path, dpi=300)
        plt.close(fig)
        del df, target, actual
        return "OK"
    except Exception:
        plt.close('all')
        return "Error"


def worker_overview_analysis(input_csv, output_dir, date_str, threshold=28.0):
    import matplotlib.pyplot as plt
    matplotlib.use('Agg')
    print(f"[OVERVIEW] Starting analysis for {date_str}...", flush=True)
    try:
        os.makedirs(output_dir, exist_ok=True)
        base_name = f"NCU_TCU_{date_str}"

        df = pd.read_csv(input_csv, dtype={"NCU": int, "TCU": int, "TC": int, "Valore": float})
        df["Timestamp"] = pd.to_datetime(df["Timestamp"], errors="coerce")
        df = df.dropna(subset=["Timestamp"])
        df = df[df["TC"].isin([4, 5])]
        all_ncus = sorted(df["NCU"].unique())

        series_by_ncu = {}
        below_thresh_rows = []

        for ncu in all_ncus:
            df_ncu = df[df["NCU"] == ncu]
            tcu_map = {}
            for tcu in df_ncu["TCU"].unique():
                part = df_ncu[df_ncu["TCU"] == tcu]
                target = part[part["TC"] == 4][["Timestamp", "Valore"]].rename(columns={"Valore": "target"})
                actual = part[part["TC"] == 5][["Timestamp", "Valore"]].rename(columns={"Valore": "actual"})
                merged = pd.merge_asof(
                    target.sort_values("Timestamp"),
                    actual.sort_values("Timestamp"),
                    on="Timestamp", direction="nearest",
                    tolerance=pd.Timedelta("1s")
                ).dropna()
                if not merged.empty:
                    tcu_map[tcu] = merged
                    low = merged[merged["actual"] < threshold]
                    for _, row in low.iterrows():
                        below_thresh_rows.append({"NCU": ncu, "TCU": tcu, "Timestamp": row["Timestamp"], "Angle": row["actual"]})
            series_by_ncu[ncu] = tcu_map

        def make_plot(ncus, fname, title):
            print(f"[OVERVIEW] Generating: {os.path.basename(fname)}", flush=True)
            fig, ax = plt.subplots(figsize=(15, 4))
            colors = {1: plt.cm.Blues, 2: plt.cm.Greens, 3: plt.cm.Reds}
            for ncu in ncus:
                tmap = series_by_ncu.get(ncu, {})
                if not tmap:
                    continue
                cmap = colors.get(ncu, plt.cm.tab10)
                tcus = sorted(tmap.keys())
                N = max(len(tcus) - 1, 1)
                for i, tcu in enumerate(tcus):
                    ax.plot(tmap[tcu]["Timestamp"], tmap[tcu]["actual"], color=cmap(i / N), lw=0.8)
            ax.axhline(threshold, color='red', linestyle='--', lw=0.8, label=f"Threshold {threshold}")
            ax.set_title(title)
            ax.grid(True, alpha=0.3)
            ax.set_ylabel("Angle (deg)")
            ax.legend(loc="upper right", fontsize=8)
            fig.tight_layout()
            fig.savefig(fname, dpi=600)
            plt.close(fig)
            plt.clf()

        make_plot(all_ncus, os.path.join(output_dir, f"{base_name}_ALL.png"), "All Trackers")
        for ncu in all_ncus:
            make_plot([ncu], os.path.join(output_dir, f"{base_name}_NCU{ncu}.png"), f"NCU {ncu} Overview")

        if below_thresh_rows:
            pd.DataFrame(below_thresh_rows).to_csv(
                os.path.join(output_dir, f"{base_name}_below_{int(threshold)}deg.csv"),
                index=False, sep=";"
            )

        del df, series_by_ncu, below_thresh_rows
        gc.collect()
        return "Overview Generation Complete"
    except Exception as e:
        print(f"[OVERVIEW] Error: {e}", flush=True)
        return f"Error: {e}"


def worker_health_check(csv_path, date_str, angle_th=28.0, dev_th=5.0):
    issues = []
    print(f"[HEALTH] Reading CSV: {csv_path}", flush=True)
    try:
        df = pd.read_csv(csv_path, dtype={"NCU": int, "TCU": int, "TC": int, "Valore": float})
        df["Timestamp"] = pd.to_datetime(df["Timestamp"])
        df = df[df["TC"].isin([4, 5])].dropna(subset=["Valore"])

        dt_obj = datetime.strptime(date_str, "%Y-%m-%d")
        sr, ss = get_sun_times(dt_obj)
        start_check = sr + timedelta(minutes=30)
        end_check = ss - timedelta(minutes=30)

        print(f"  Sunrise: {sr.strftime('%H:%M')}  Sunset: {ss.strftime('%H:%M')}", flush=True)

        groups = list(df.groupby(["NCU", "TCU"]))
        print(f"[HEALTH] Analyzing {len(groups)} trackers...", flush=True)

        for i, ((ncu, tcu), g) in enumerate(groups):
            if i % 50 == 0:
                print(f"[HEALTH] {i}/{len(groups)} analyzed...", flush=True)

            actual_df = g[g["TC"] == 5].sort_values("Timestamp")
            target_df = g[g["TC"] == 4].sort_values("Timestamp")

            if actual_df.empty:
                issues.append({"NCU": ncu, "TCU": tcu, "Type": "NO DATA",
                               "Msg": "No actual angle data", "Sev": "High"})
                continue

            # COMM ERROR: values outside 30-150 range
            if (actual_df["Valore"] <= 30).any() or (actual_df["Valore"] >= 150).any():
                issues.append({"NCU": ncu, "TCU": tcu, "Type": "COMM ERROR",
                               "Msg": "Value outside 30-150 range", "Sev": "Communication Error"})

            # SUDDEN JUMP
            diff_series = actual_df["Valore"].diff().abs()
            prev_vals = actual_df["Valore"].shift(1)
            valid_jumps = actual_df[(diff_series > 15) & (actual_df["Valore"] > 30) & (prev_vals > 30)]
            if not valid_jumps.empty:
                issues.append({"NCU": ncu, "TCU": tcu, "Type": "SUDDEN JUMP",
                               "Msg": f"Max jump {diff_series.max():.1f} deg", "Sev": "Medium"})

            # STUCK / HIGH WIND (daylight hours only)
            daylight = actual_df[
                (actual_df["Timestamp"] >= start_check) & (actual_df["Timestamp"] <= end_check)
            ].copy()

            if len(daylight) > 15:
                daylight_idx = daylight.set_index("Timestamp")
                rolling_std = daylight_idx["Valore"].rolling("15min").std()
                stuck_timestamps = rolling_std[rolling_std == 0].index

                if not stuck_timestamps.empty:
                    stuck_period = daylight[daylight["Timestamp"].isin(stuck_timestamps)]
                    merged = pd.merge_asof(
                        stuck_period.sort_values("Timestamp"),
                        target_df.sort_values("Timestamp"),
                        on="Timestamp", direction="nearest",
                        tolerance=pd.Timedelta("2min"),
                        suffixes=("_act", "_tgt")
                    )
                    if "Valore_act" in merged.columns and "Valore_tgt" in merged.columns:
                        merged["deviation"] = (merged["Valore_act"] - merged["Valore_tgt"]).abs()
                        problematic = merged[merged["deviation"] > dev_th]
                        if not problematic.empty:
                            stuck_val = problematic["Valore_act"].mean()
                            # Determine stuck time range
                            t_start = problematic["Timestamp"].min()
                            t_end = problematic["Timestamp"].max()
                            duration_min = (t_end - t_start).total_seconds() / 60
                            if 85 <= stuck_val <= 98:
                                issues.append({
                                    "NCU": ncu, "TCU": tcu, "Type": "HIGH WIND MODE",
                                    "Msg": f"Stowed at {stuck_val:.1f} deg | Dev>{dev_th} | {duration_min:.0f} min",
                                    "Sev": "Low",
                                    "StuckAngle": stuck_val, "StuckStart": t_start,
                                    "StuckEnd": t_end, "AvgDev": problematic["deviation"].mean()
                                })
                            else:
                                issues.append({
                                    "NCU": ncu, "TCU": tcu, "Type": "STUCK",
                                    "Msg": f"Stuck at {stuck_val:.1f} deg | Dev>{dev_th} | {duration_min:.0f} min",
                                    "Sev": "Medium",
                                    "StuckAngle": stuck_val, "StuckStart": t_start,
                                    "StuckEnd": t_end, "AvgDev": problematic["deviation"].mean()
                                })

            # LOW ANGLE
            min_val = actual_df["Valore"].min()
            if 0.1 < min_val < 10:
                issues.append({"NCU": ncu, "TCU": tcu, "Type": "LOW ANGLE",
                               "Msg": f"Min angle {min_val:.1f} deg", "Sev": "Low"})

        del df, groups
        gc.collect()
        return issues

    except Exception as e:
        print(f"[HEALTH] Error: {e}", flush=True)
        return [{"NCU": 0, "TCU": 0, "Type": "ERROR", "Msg": str(e), "Sev": "Critical"}]


def worker_generate_pdf_from_images(overview_dir, plots_dir, output_pdf, date_str):
    """Button 6: Full PDF using saved PNG images (all 370 trackers)."""
    import matplotlib.pyplot as plt
    import matplotlib.image as mpimg
    from matplotlib.backends.backend_pdf import PdfPages
    matplotlib.use('Agg')

    print(f"[PDF] Building full PDF from saved images...", flush=True)
    try:
        os.makedirs(os.path.dirname(output_pdf), exist_ok=True)
        with PdfPages(output_pdf) as pdf:

            # --- Overview pages (ALL, NCU1, NCU2, NCU3) ---
            for sfx in ["ALL", "NCU1", "NCU2", "NCU3"]:
                img_path = os.path.join(overview_dir, f"NCU_TCU_{date_str}_{sfx}.png")
                if os.path.exists(img_path):
                    print(f"  -> Overview: {sfx}", flush=True)
                    img = mpimg.imread(img_path)
                    fig, ax = plt.subplots(figsize=(15.5, 11.2))
                    ax.imshow(img)
                    ax.axis('off')
                    pdf.savefig(fig, orientation='landscape', bbox_inches='tight', dpi=150)
                    plt.close(fig)
                    del img

            # --- Individual tracker pages from saved PNGs ---
            img_files = sorted(
                glob.glob(os.path.join(plots_dir, "TX_*.png")),
                key=lambda p: (
                    int(re.search(r"TX_(\d+)_", p).group(1)) if re.search(r"TX_(\d+)_", p) else 0,
                    int(re.search(r"TCU_(\d+)", p).group(1)) if re.search(r"TCU_(\d+)", p) else 0
                )
            )

            total = len(img_files)
            print(f"[PDF] Adding {total} tracker images...", flush=True)
            for count, img_path in enumerate(img_files, 1):
                if count % 50 == 0:
                    print(f"[PDF] {count}/{total} pages done...", flush=True)
                    gc.collect()
                try:
                    img = mpimg.imread(img_path)
                    fig, ax = plt.subplots(figsize=(10, 6))
                    ax.imshow(img)
                    ax.axis('off')
                    pdf.savefig(fig, bbox_inches='tight', dpi=150)
                    plt.close(fig)
                    del img
                except Exception as e:
                    print(f"  [WARNING] Could not add {os.path.basename(img_path)}: {e}", flush=True)

        return f"Full PDF saved to:\n{output_pdf}"
    except Exception as e:
        return f"Error generating PDF: {e}"


def worker_generate_random_pdf(csv_path, img_folder, output_pdf, date_str):
    """Button 7: Random sample PDF (5 trackers per NCU) using saved images."""
    import matplotlib.pyplot as plt
    import matplotlib.image as mpimg
    from matplotlib.backends.backend_pdf import PdfPages
    import random
    matplotlib.use('Agg')

    print(f"[RANDOM PDF] Building random sample PDF...", flush=True)
    try:
        os.makedirs(os.path.dirname(output_pdf), exist_ok=True)
        df = pd.read_csv(csv_path, dtype={"NCU": int, "TCU": int, "TC": int, "Valore": float})
        df["Timestamp"] = pd.to_datetime(df["Timestamp"])
        df = df.dropna(subset=["Valore"])
        overview_dir = os.path.dirname(img_folder)

        with PdfPages(output_pdf) as pdf:

            # --- Overview pages (use saved high-quality PNGs) ---
            for sfx in ["ALL", "NCU1", "NCU2", "NCU3"]:
                img_path = os.path.join(overview_dir, f"NCU_TCU_{date_str}_{sfx}.png")
                if os.path.exists(img_path):
                    print(f"  -> Overview: {sfx}", flush=True)
                    img = mpimg.imread(img_path)
                    fig, ax = plt.subplots(figsize=(15.5, 11.2))
                    ax.imshow(img)
                    ax.axis('off')
                    pdf.savefig(fig, orientation='landscape', bbox_inches='tight', dpi=150)
                    plt.close(fig)
                    del img

            # --- Individual: 5 random TCUs per NCU ---
            for ncu, ncu_data in df.groupby("NCU"):
                tcus = list(ncu_data["TCU"].unique())
                selected = sorted(random.sample(tcus, k=min(len(tcus), 5)))
                print(f"  -> NCU {int(ncu)}: {selected}", flush=True)

                for tcu in selected:
                    img_name = f"TX_{int(ncu)}_TCU_{int(tcu)}.png"
                    img_path = os.path.join(img_folder, img_name)

                    fig, ax = plt.subplots(figsize=(10, 6))
                    if os.path.exists(img_path):
                        try:
                            img = mpimg.imread(img_path)
                            ax.imshow(img)
                            ax.axis('off')
                            del img
                        except Exception as e:
                            print(f"    Error loading {img_name}: {e}", flush=True)
                    else:
                        group = ncu_data[ncu_data["TCU"] == tcu]
                        target = group[group["TC"] == 4].sort_values("Timestamp")
                        actual = group[group["TC"] == 5].sort_values("Timestamp")
                        if not target.empty:
                            ax.plot(target["Timestamp"], target["Valore"], label="Target", color="royalblue", lw=1)
                        if not actual.empty:
                            ax.plot(actual["Timestamp"], actual["Valore"], label="Actual", color="tomato", lw=1)
                        ax.set_title(f"NCU {int(ncu)} - TCU {int(tcu)}")
                        ax.set_ylabel("Angle (deg)")
                        ax.grid(True, alpha=0.3)
                        ax.legend(loc="upper right")
                    pdf.savefig(fig, dpi=150, bbox_inches='tight')
                    plt.close(fig)

                del ncu_data
                gc.collect()

        del df
        gc.collect()
        return f"Random PDF saved to:\n{output_pdf}"
    except Exception as e:
        return f"Error: {e}"


def worker_generate_multiday_plots(base, start_date, end_date, selected_ncus, output_dir):
    """Advanced: Generate multi-day NCU overview plots from existing merged CSVs."""
    import matplotlib.pyplot as plt
    matplotlib.use('Agg')

    print(f"[MULTIDAY] Generating plots {start_date} to {end_date} | NCUs: {selected_ncus}", flush=True)
    os.makedirs(output_dir, exist_ok=True)

    # Collect dates
    dates = []
    cur = datetime.strptime(start_date, "%Y-%m-%d")
    end = datetime.strptime(end_date, "%Y-%m-%d")
    while cur <= end:
        dates.append(cur.strftime("%Y-%m-%d"))
        cur += timedelta(days=1)

    all_dfs = []
    for d in dates:
        y, m = d[:4], d[5:7]
        csv_path = os.path.join(base, "03_Merged_files", y, m, f"{d}_1min_merged.csv")
        if os.path.exists(csv_path):
            try:
                df = pd.read_csv(csv_path, dtype={"NCU": int, "TCU": int, "TC": int, "Valore": float})
                df["Timestamp"] = pd.to_datetime(df["Timestamp"])
                all_dfs.append(df)
                print(f"[MULTIDAY] Loaded {d}", flush=True)
            except Exception as e:
                print(f"[MULTIDAY] Skipped {d}: {e}", flush=True)
        else:
            print(f"[MULTIDAY] Not found: {d}", flush=True)

    if not all_dfs:
        return "No CSV data found for the selected date range."

    combined = pd.concat(all_dfs, ignore_index=True)
    combined = combined[combined["TC"] == 5]
    date_range_str = f"{start_date}_to_{end_date}"

    saved = []
    for ncu in selected_ncus:
        ncu_data = combined[combined["NCU"] == ncu]
        if ncu_data.empty:
            continue
        fig, ax = plt.subplots(figsize=(15, 5))
        for tcu in sorted(ncu_data["TCU"].unique()):
            t = ncu_data[ncu_data["TCU"] == tcu]
            ax.plot(t["Timestamp"], t["Valore"], lw=0.5, alpha=0.7)
        ax.set_title(f"NCU {ncu} - Multi-day Overview ({start_date} to {end_date})")
        ax.set_ylabel("Angle (deg)")
        ax.grid(True, alpha=0.3)
        fig.tight_layout()
        fname = os.path.join(output_dir, f"Multiday_{date_range_str}_NCU{ncu}.png")
        fig.savefig(fname, dpi=600, bbox_inches="tight")
        plt.close(fig)
        saved.append(fname)
        print(f"[MULTIDAY] Saved NCU {ncu}: {fname}", flush=True)

    del combined, all_dfs
    gc.collect()

    if saved:
        # Open the output folder
        os.startfile(os.path.normpath(output_dir))
        return f"Multi-day plots saved ({len(saved)} files):\n{output_dir}"
    return "No plots were generated (check NCU selection and data availability)."


def worker_generate_official_report(base, date_str, template_path, output_path, sample_n=0):
    """
    Generate official DOCX report matching the PDF template style.
    sample_n=0 -> full report (all trackers)
    sample_n>0 -> sample report (sample_n random TCUs per NCU)

    Page layout:
      Front matter : A4 portrait  (21.0 x 29.7 cm)
      NCU overviews: Wide landscape (42.0 x 15.0 cm)  - Allegato A
      Tracker plots : A4 landscape  (29.7 x 21.0 cm)  - Allegato B
    """
    try:
        from docx import Document
        from docx.shared import Cm
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        import random as rnd
    except ImportError:
        return "Error: python-docx not installed. Run: pip install python-docx"

    # ---- helpers ----
    def _tw(cm_val):
        """Centimetres to twips (Word internal unit: 1 inch = 1440 twips)."""
        return int(round(cm_val * 1440 / 2.54))

    def _set_para_section(para, w_cm, h_cm, mt=1.5, mb=1.5, ml=1.5, mr=1.5):
        """
        Embed a <w:sectPr> in this paragraph's <w:pPr>.
        In Word semantics this CLOSES the current section: all paragraphs from
        the previous sectPr up to and including this one use w_cm x h_cm.
        The NEXT section starts from the following paragraph.
        """
        pPr = para._element.get_or_add_pPr()
        for old in pPr.findall(qn('w:sectPr')):
            pPr.remove(old)
        s = OxmlElement('w:sectPr')
        pgSz = OxmlElement('w:pgSz')
        pgSz.set(qn('w:w'), str(_tw(w_cm)))
        pgSz.set(qn('w:h'), str(_tw(h_cm)))
        if w_cm > h_cm:
            pgSz.set(qn('w:orient'), 'landscape')
        s.append(pgSz)
        pgMar = OxmlElement('w:pgMar')
        pgMar.set(qn('w:top'),    str(_tw(mt)))
        pgMar.set(qn('w:right'),  str(_tw(mr)))
        pgMar.set(qn('w:bottom'), str(_tw(mb)))
        pgMar.set(qn('w:left'),   str(_tw(ml)))
        pgMar.set(qn('w:header'), str(_tw(0.5)))
        pgMar.set(qn('w:footer'), str(_tw(0.5)))
        pgMar.set(qn('w:gutter'), '0')
        s.append(pgMar)
        pPr.append(s)

    def _set_body_section(w_cm, h_cm, mt=1.5, mb=1.5, ml=1.5, mr=1.5):
        """Update the body-level <w:sectPr> (governs the final / only remaining section)."""
        body = doc.element.body
        s = body.find(qn('w:sectPr'))
        if s is None:
            s = OxmlElement('w:sectPr')
            body.append(s)
        pgSz = s.find(qn('w:pgSz'))
        if pgSz is None:
            pgSz = OxmlElement('w:pgSz')
            s.insert(0, pgSz)
        pgSz.set(qn('w:w'), str(_tw(w_cm)))
        pgSz.set(qn('w:h'), str(_tw(h_cm)))
        if w_cm > h_cm:
            pgSz.set(qn('w:orient'), 'landscape')
        else:
            pgSz.attrib.pop(qn('w:orient'), None)
        pgMar = s.find(qn('w:pgMar'))
        if pgMar is None:
            pgMar = OxmlElement('w:pgMar')
            s.append(pgMar)
        pgMar.set(qn('w:top'),    str(_tw(mt)))
        pgMar.set(qn('w:right'),  str(_tw(mr)))
        pgMar.set(qn('w:bottom'), str(_tw(mb)))
        pgMar.set(qn('w:left'),   str(_tw(ml)))

    def _insert_image(anchor, img_path, w_cm):
        """Add an image paragraph immediately after anchor. Returns the new paragraph."""
        ip = doc.add_paragraph()
        ip.alignment = 1
        run = ip.add_run()
        try:
            run.add_picture(img_path, width=Cm(w_cm))
        except Exception:
            run.text = f"[Image unavailable: {os.path.basename(img_path)}]"
        anchor._element.addnext(ip._element)
        return ip

    # ---- page geometry ----
    # NCU overview: very wide narrow landscape (mimics the PDF template 38.1x10.2cm but taller for margins)
    NCU_W, NCU_H   = 42.0, 15.0   # cm  (A3-width wide landscape)
    NCU_IMG_W      = 38.5          # cm  image width inside NCU page
    NCU_MARGIN     = (1.0, 1.75, 1.0, 1.75)  # mt, mb, ml, mr

    # Tracker: A4 landscape
    TRK_W, TRK_H   = 29.7, 21.0   # cm
    TRK_IMG_W      = 26.0          # cm  image width inside A4 landscape
    TRK_MARGIN     = (1.5, 1.5, 1.5, 1.5)

    print(f"[REPORT] Building {'sample' if sample_n else 'full'} report for {date_str}...", flush=True)

    try:
        dt = datetime.strptime(date_str, "%Y-%m-%d")
        dd, mm_s, yyyy = dt.strftime("%d"), dt.strftime("%m"), dt.strftime("%Y")
        mmmm = ITALIAN_MONTHS[dt.month]
        mese_str = f"{mmmm} {yyyy}"

        replacements = {
            "MESE 2026": mese_str, "MESE 2025": mese_str,
            "DD/MM/YYYY": f"{dd}/{mm_s}/{yyyy}",
            "(DD MMMM 2026)": f"({dd} {mmmm} {yyyy})",
            "(DD MMMM 2025)": f"({dd} {mmmm} {yyyy})",
            "DD MMMM YYYY": f"{dd} {mmmm} {yyyy}",
            "MMMM": mmmm, "YYYY": yyyy, "DD": dd, "MM": mm_s,
        }

        shutil.copy2(template_path, output_path)
        doc = Document(output_path)

        def repl_para(para):
            if not para.runs:
                return
            full = "".join(r.text for r in para.runs)
            new = full
            for old, nv in replacements.items():
                new = new.replace(old, nv)
            if new != full:
                para.runs[0].text = new
                for r in para.runs[1:]:
                    r.text = ""

        for para in doc.paragraphs:
            repl_para(para)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        repl_para(para)

        allegato_a_anchor = next((p for p in doc.paragraphs if p.text.strip() == "Riepilogo complessivo"), None)
        allegato_b_anchor = next((p for p in doc.paragraphs if p.text.strip() == "Riepilogo di ogni tracker"), None)

        y, m = date_str[:4], date_str[5:7]
        overview_dir = os.path.join(base, "04_Tracker_plots_angles", y, m, date_str)
        plots_dir    = os.path.join(overview_dir, "each_tracker_plots")

        # ---- Allegato A: NCU overview images (wide landscape section) ----
        if allegato_a_anchor:
            # This sectPr closes the A4-portrait section at "Riepilogo complessivo"
            _set_para_section(allegato_a_anchor, 21.0, 29.7, mt=2.0, mb=2.0, ml=2.5, mr=2.0)

            last = allegato_a_anchor
            ncu_added = 0
            for sfx in ["ALL", "NCU1", "NCU2", "NCU3"]:
                img = os.path.join(overview_dir, f"NCU_TCU_{date_str}_{sfx}.png")
                if os.path.exists(img):
                    last = _insert_image(last, img, NCU_IMG_W)
                    ncu_added += 1
                    print(f"[REPORT] Added NCU overview: {sfx}", flush=True)
                else:
                    print(f"[REPORT] Missing NCU overview: {sfx}", flush=True)

            # Close the wide-landscape section at the last NCU image
            if ncu_added:
                _set_para_section(last, NCU_W, NCU_H,
                                  mt=NCU_MARGIN[0], mb=NCU_MARGIN[1],
                                  ml=NCU_MARGIN[2], mr=NCU_MARGIN[3])
        else:
            print("[REPORT] Warning: 'Riepilogo complessivo' not found.", flush=True)

        # ---- Allegato B: tracker images (A4 landscape section) ----
        if allegato_b_anchor:
            def _sort_key(path):
                n = re.search(r"TX_(\d+)_TCU_(\d+)", os.path.basename(path))
                return (int(n.group(1)), int(n.group(2))) if n else (999, 999)

            all_files = sorted(glob.glob(os.path.join(plots_dir, "TX_*.png")), key=_sort_key)

            if sample_n > 0:
                by_ncu = {}
                for f in all_files:
                    n = re.search(r"TX_(\d+)_TCU_", os.path.basename(f))
                    if n:
                        by_ncu.setdefault(int(n.group(1)), []).append(f)
                img_files = []
                for ncu_id in sorted(by_ncu):
                    pool = by_ncu[ncu_id]
                    img_files.extend(sorted(rnd.sample(pool, min(sample_n, len(pool))), key=_sort_key))
            else:
                img_files = all_files

            total = len(img_files)
            print(f"[REPORT] Adding {total} tracker pages (A4 landscape)...", flush=True)
            last = allegato_b_anchor

            for idx, img_path in enumerate(img_files):
                if idx % 50 == 0:
                    print(f"[REPORT] {idx}/{total}...", flush=True)
                last = _insert_image(last, img_path, TRK_IMG_W)
        else:
            print("[REPORT] Warning: 'Riepilogo di ogni tracker' not found.", flush=True)

        # Body sectPr: A4 landscape governs all tracker pages (the final section)
        _set_body_section(TRK_W, TRK_H,
                          mt=TRK_MARGIN[0], mb=TRK_MARGIN[1],
                          ml=TRK_MARGIN[2], mr=TRK_MARGIN[3])

        doc.save(output_path)
        label = "Sample" if sample_n else "Full"
        return f"{label} official report saved to:\n{output_path}"

    except Exception as e:
        import traceback
        return f"Error generating report:\n{e}\n{traceback.format_exc()}"


def check_missing_hours(base, date_str):
    """Check source files and merged CSV for missing hourly data. Returns list of warning strings."""
    y, m = date_str[:4], date_str[5:7]
    warnings = []

    def format_gaps(missing):
        if not missing:
            return ""
        parts = []
        start = prev = missing[0]
        for h in missing[1:]:
            if h != prev + 1:
                parts.append(f"{start:02d}:00" if start == prev else f"{start:02d}:00-{prev + 1:02d}:00")
                start = h
            prev = h
        parts.append(f"{start:02d}:00" if start == prev else f"{start:02d}:00-{prev + 1:02d}:00")
        return ", ".join(parts)

    # Check 1: source files
    src_dir = os.path.join(base, "01_Original_files", y, m, date_str)
    if os.path.exists(src_dir):
        files = sorted(glob.glob(os.path.join(src_dir, "*.csv")) + glob.glob(os.path.join(src_dir, "*.txt")))
        if files:
            hours_found = set()
            for f in files:
                for enc in ["utf-8-sig", "utf-16", "latin1"]:
                    try:
                        with open(f, "r", encoding=enc, errors="ignore") as fh:
                            for line in fh:
                                hit = re.search(r"\d{2}/\d{2}/\d{4}\s+(\d{2}):", line)
                                if hit:
                                    hours_found.add(int(hit.group(1)))
                                    break
                        break
                    except UnicodeDecodeError:
                        continue
            if hours_found:
                missing = sorted(set(range(24)) - hours_found)
                if missing:
                    warnings.append(f"SOURCE FILES: Missing hours {format_gaps(missing)}")
            else:
                warnings.append(f"SOURCE FILES: {len(files)} files found but timestamps not parseable")
        else:
            warnings.append("SOURCE FILES: No CSV/TXT files found in source folder")
    else:
        warnings.append(f"SOURCE FILES: Folder not found:\n{src_dir}")

    # Check 2: merged CSV
    merged = os.path.join(base, "03_Merged_files", y, m, f"{date_str}_1min_merged.csv")
    if os.path.exists(merged):
        try:
            df = pd.read_csv(merged, usecols=["Timestamp"])
            df["Timestamp"] = pd.to_datetime(df["Timestamp"], errors="coerce").dropna()
            if not df.empty:
                present = set(df["Timestamp"].dt.hour.unique())
                missing = sorted(set(range(24)) - present)
                if missing:
                    warnings.append(f"MERGED DATA: Missing hours {format_gaps(missing)}")
                else:
                    warnings.append("MERGED DATA: All 24 hours present")
        except Exception as e:
            warnings.append(f"MERGED DATA: Could not read ({e})")
    else:
        warnings.append("MERGED DATA: Merged CSV not found (run Extract + Merge first)")

    return warnings


# ==============================================================================
# CALENDAR POPUP
# ==============================================================================

class CalendarPopup(tk.Toplevel):
    """Popup calendar picker styled with Material Design colors."""

    def __init__(self, parent, initial_year, initial_month, initial_day, callback):
        super().__init__(parent)
        self.title("Select Date")
        self.resizable(False, False)
        self.callback = callback
        self.transient(parent)

        if HAS_TKCALENDAR:
            cal = TkCalendar(
                self, selectmode="day",
                year=initial_year, month=initial_month, day=initial_day,
                background="#1a73e8", foreground="white",
                selectbackground="#1a73e8", selectforeground="white",
                headersbackground="#1565c0", headersforeground="white",
                normalbackground="white", normalforeground="#202124",
                weekendbackground="#f8f9fa", weekendforeground="#5f6368",
                othermonthbackground="#f1f3f4", othermonthforeground="#bdc1c6",
                bordercolor="#dadce0", font=("Segoe UI", 10),
                showweeknumbers=False, cursor="hand2"
            )
            cal.pack(padx=12, pady=12)
            btn_frame = tk.Frame(self, bg="white")
            btn_frame.pack(fill="x", padx=12, pady=(0, 12))
            tk.Button(
                btn_frame, text="Cancel", bg="#f1f3f4", fg="#5f6368",
                relief="flat", padx=12, pady=6, cursor="hand2",
                command=self.destroy
            ).pack(side="right", padx=(4, 0))
            tk.Button(
                btn_frame, text="Select", bg="#1a73e8", fg="white",
                relief="flat", padx=12, pady=6, cursor="hand2",
                command=lambda: self._select(cal)
            ).pack(side="right")
        else:
            tk.Label(self, text="tkcalendar not installed.\nRun: pip install tkcalendar").pack(padx=20, pady=20)
            ttk.Button(self, text="Close", command=self.destroy).pack(pady=8)
            return

        self.grab_set()
        self.update_idletasks()
        pw = parent.winfo_rootx() + parent.winfo_width() // 2
        ph = parent.winfo_rooty() + parent.winfo_height() // 2
        self.geometry(f"+{pw - 130}+{ph - 130}")

    def _select(self, cal):
        date_str = cal.get_date()
        for fmt in ["%m/%d/%y", "%d/%m/%Y", "%Y-%m-%d"]:
            try:
                dt = datetime.strptime(date_str, fmt)
                self.callback(dt)
                self.destroy()
                return
            except ValueError:
                continue
        messagebox.showerror("Error", f"Could not parse date: {date_str}")
        self.destroy()


# ==============================================================================
# ADVANCED SETTINGS DIALOG
# ==============================================================================

class AdvancedSettingsDialog(tk.Toplevel):

    def __init__(self, parent, app):
        super().__init__(parent)
        self.title("Advanced Settings")
        self.geometry("520x420")
        self.resizable(False, False)
        self.transient(parent)
        self.app = app

        nb = ttk.Notebook(self)
        nb.pack(fill="both", expand=True, padx=10, pady=10)

        # --- Tab 1: Health Check ---
        tab_hc = ttk.Frame(nb)
        nb.add(tab_hc, text="Health Check")
        ttk.Label(tab_hc, text="Min angle threshold (deg):").grid(row=0, column=0, sticky="w", padx=15, pady=8)
        self.angle_th = tk.StringVar(value=str(app.angle_threshold))
        ttk.Entry(tab_hc, textvariable=self.angle_th, width=8).grid(row=0, column=1, padx=5)
        ttk.Label(tab_hc, text="Deviation threshold (deg):").grid(row=1, column=0, sticky="w", padx=15, pady=8)
        self.dev_th = tk.StringVar(value=str(app.dev_threshold))
        ttk.Entry(tab_hc, textvariable=self.dev_th, width=8).grid(row=1, column=1, padx=5)
        ttk.Label(tab_hc, text="(Stuck: tracker not moving when target is. High wind: stuck at 85-98 deg)",
                  foreground="#5f6368", wraplength=380).grid(row=2, column=0, columnspan=2, sticky="w", padx=15, pady=4)

        # --- Tab 2: Multi-day Plots ---
        tab_md = ttk.Frame(nb)
        nb.add(tab_md, text="Multi-day Plot")

        ttk.Label(tab_md, text="Start date (YYYY-MM-DD):").grid(row=0, column=0, sticky="w", padx=15, pady=8)
        self.md_start = tk.StringVar(value=app.date_val.get())
        ttk.Entry(tab_md, textvariable=self.md_start, width=14).grid(row=0, column=1, padx=5)

        ttk.Label(tab_md, text="End date (YYYY-MM-DD):").grid(row=1, column=0, sticky="w", padx=15, pady=8)
        self.md_end = tk.StringVar(value=app.date_val.get())
        ttk.Entry(tab_md, textvariable=self.md_end, width=14).grid(row=1, column=1, padx=5)

        ttk.Label(tab_md, text="Include NCUs:").grid(row=2, column=0, sticky="w", padx=15, pady=8)
        ncu_frame = ttk.Frame(tab_md)
        ncu_frame.grid(row=2, column=1, sticky="w")
        self.ncu1_var = tk.BooleanVar(value=True)
        self.ncu2_var = tk.BooleanVar(value=True)
        self.ncu3_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(ncu_frame, text="NCU1", variable=self.ncu1_var).pack(side="left")
        ttk.Checkbutton(ncu_frame, text="NCU2", variable=self.ncu2_var).pack(side="left")
        ttk.Checkbutton(ncu_frame, text="NCU3", variable=self.ncu3_var).pack(side="left")

        ttk.Label(tab_md, text="Output saved to 04_Tracker_plots_angles/multiday/",
                  foreground="#5f6368").grid(row=3, column=0, columnspan=2, sticky="w", padx=15, pady=4)

        tk.Button(
            tab_md, text="Generate Multi-day Plots", bg="#1a73e8", fg="white",
            relief="flat", padx=10, pady=6, cursor="hand2",
            command=self._run_multiday
        ).grid(row=4, column=0, columnspan=2, pady=15)

        # --- Tab 3: Missing Hours Check ---
        tab_hr = ttk.Frame(nb)
        nb.add(tab_hr, text="Hour Check")
        ttk.Label(tab_hr, text="Check for missing hourly data files\nand gaps in the merged CSV.",
                  justify="left").pack(padx=15, pady=15)
        tk.Button(
            tab_hr, text="Check Missing Hours", bg="#34a853", fg="white",
            relief="flat", padx=10, pady=6, cursor="hand2",
            command=self._run_hour_check
        ).pack(pady=5)
        self.hr_result = tk.Text(tab_hr, height=8, state="disabled", bg="#f8f9fa", wrap="word")
        self.hr_result.pack(fill="both", expand=True, padx=10, pady=5)

        # --- Bottom buttons ---
        btn_frame = ttk.Frame(self)
        btn_frame.pack(fill="x", padx=10, pady=(0, 10))
        ttk.Button(btn_frame, text="Apply & Close", command=self._apply).pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Cancel", command=self.destroy).pack(side="right")

        self.grab_set()

    def _apply(self):
        try:
            self.app.angle_threshold = float(self.angle_th.get())
            self.app.dev_threshold = float(self.dev_th.get())
        except ValueError:
            messagebox.showerror("Invalid", "Thresholds must be numbers.", parent=self)
            return
        self.destroy()

    def _run_multiday(self):
        ncus = []
        if self.ncu1_var.get():
            ncus.append(1)
        if self.ncu2_var.get():
            ncus.append(2)
        if self.ncu3_var.get():
            ncus.append(3)
        if not ncus:
            messagebox.showwarning("No NCU", "Select at least one NCU.", parent=self)
            return
        start = self.md_start.get().strip()
        end = self.md_end.get().strip()
        try:
            datetime.strptime(start, "%Y-%m-%d")
            datetime.strptime(end, "%Y-%m-%d")
        except ValueError:
            messagebox.showerror("Invalid date", "Dates must be YYYY-MM-DD format.", parent=self)
            return
        base = self.app.folder_path.get()
        out_dir = os.path.join(base, "04_Tracker_plots_angles", "multiday", f"{start}_to_{end}")
        print(f"\n[MULTIDAY] Starting: {start} to {end} | NCUs: {ncus}")

        def run():
            msg = worker_generate_multiday_plots(base, start, end, ncus, out_dir)
            self.app.root.after(0, lambda: messagebox.showinfo("Multi-day Plot", msg, parent=self.app.root))

        threading.Thread(target=run, daemon=True).start()

    def _run_hour_check(self):
        base = self.app.folder_path.get()
        date = self.app.date_val.get()
        warnings = check_missing_hours(base, date)
        self.hr_result.config(state="normal")
        self.hr_result.delete("1.0", "end")
        if not warnings:
            self.hr_result.insert("end", "No issues found - all hours present.")
        else:
            for w in warnings:
                self.hr_result.insert("end", w + "\n\n")
        self.hr_result.config(state="disabled")


# ==============================================================================
# TEXT REDIRECTOR
# ==============================================================================

class TextRedirector:
    def __init__(self, widget):
        self.widget = widget
        self.terminal = sys.__stdout__

    def write(self, s):
        try:
            self.widget.config(state="normal")
            self.widget.insert("end", s)
            self.widget.see("end")
            self.widget.config(state="disabled")
            self.widget.update_idletasks()
        except Exception:
            pass
        if self.terminal:
            self.terminal.write(s)
            self.terminal.flush()

    def flush(self):
        if self.terminal:
            self.terminal.flush()


# ==============================================================================
# MAIN APPLICATION
# ==============================================================================

class TrackerSuiteApp:

    def __init__(self, root):
        self.root = root
        self.root.title("GET - SCADA Tracker Suite v2.12")
        self.root.geometry("900x820")
        self.root.configure(bg="#f1f3f4")

        # --- Logo ---
        if hasattr(sys, '_MEIPASS'):
            self.logo_path = os.path.join(sys._MEIPASS, "logo.png")
        else:
            self.logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logo.png")

        self.cpu_cores = min(multiprocessing.cpu_count(), 8)
        self.folder_path = tk.StringVar(value=DEFAULT_ROOT)
        self.date_val = tk.StringVar()
        today = datetime.now()
        self.var_year = tk.StringVar(value=str(today.year))
        self.var_month = tk.StringVar(value=str(today.month).zfill(2))
        self.var_day = tk.StringVar(value=str(today.day).zfill(2))
        self.update_date_str()

        self.angle_threshold = 28.0
        self.dev_threshold = 5.0
        self.health_issues = {}
        self._progress_max = 0

        self._build_ui()
        sys.stdout = TextRedirector(self.log_text)
        print("=== GET - SCADA Tracker Suite v2.12 ===")

        self.var_year.trace_add("write", self.update_date_str)
        self.var_month.trace_add("write", self.update_date_str)
        self.var_day.trace_add("write", self.update_date_str)
        self.folder_path.trace_add("write", lambda *a: self.root.after(500, self.check_status))
        self.date_val.trace_add("write", lambda *a: self.root.after(300, self.check_status))

    def update_date_str(self, *args):
        try:
            y = self.var_year.get()
            m = self.var_month.get().zfill(2)
            d = self.var_day.get().zfill(2)
            self.date_val.set(f"{y}-{m}-{d}")
        except Exception:
            pass

    def _build_ui(self):
        # --- Logo bar ---
        logo_bar = tk.Frame(self.root, bg="white", height=70)
        logo_bar.pack(fill="x")
        if os.path.exists(self.logo_path):
            try:
                img = Image.open(self.logo_path)
                aspect = img.width / img.height
                img = img.resize((int(50 * aspect), 50), Image.LANCZOS)
                self._logo_img = ImageTk.PhotoImage(img)
                tk.Label(logo_bar, image=self._logo_img, bg="white").pack(side="left", padx=15, pady=10)
            except Exception:
                pass
        tk.Label(logo_bar, text="SCADA Tracker Suite v2.12", bg="white",
                 font=("Segoe UI", 14, "bold"), fg="#1a73e8").pack(side="left", padx=10)

        # --- Config frame ---
        cfg = ttk.LabelFrame(self.root, text="Configuration", padding=8)
        cfg.pack(fill="x", padx=10, pady=5)

        row1 = ttk.Frame(cfg)
        row1.pack(fill="x", pady=2)
        ttk.Label(row1, text="Root Folder:").pack(side="left")
        ttk.Entry(row1, textvariable=self.folder_path, width=62).pack(side="left", padx=5)
        ttk.Button(row1, text="Browse", command=self.browse).pack(side="left")

        row2 = ttk.Frame(cfg)
        row2.pack(fill="x", pady=4)
        ttk.Label(row2, text="Date:").pack(side="left")
        ttk.Spinbox(row2, from_=2020, to=2035, textvariable=self.var_year, width=6).pack(side="left", padx=(5, 2))
        ttk.Label(row2, text="-").pack(side="left")
        ttk.Spinbox(row2, from_=1, to=12, textvariable=self.var_month, width=3, format="%02.0f").pack(side="left")
        ttk.Label(row2, text="-").pack(side="left")
        ttk.Spinbox(row2, from_=1, to=31, textvariable=self.var_day, width=3, format="%02.0f").pack(side="left")
        # Calendar popup button
        tk.Button(row2, text="[cal]", bg="#e8f0fe", fg="#1a73e8", relief="flat",
                  padx=6, cursor="hand2", font=("Segoe UI", 9),
                  command=self.open_calendar).pack(side="left", padx=8)
        ttk.Label(row2, textvariable=self.date_val, foreground="#5f6368", font=("Segoe UI", 9)).pack(side="left", padx=5)
        ttk.Button(row2, text="Refresh Status", command=self.check_status).pack(side="left", padx=15)

        # --- Tabs ---
        self.tabs = ttk.Notebook(self.root)
        self.tabs.pack(fill="both", expand=True, padx=10, pady=5)
        self.tab_pipe = ttk.Frame(self.tabs)
        self.tab_dash = ttk.Frame(self.tabs)
        self.tabs.add(self.tab_pipe, text="Pipeline")
        self.tabs.add(self.tab_dash, text="Health Dashboard")

        self._build_pipeline_tab()
        self._build_dashboard_tab()

    def _build_pipeline_tab(self):
        outer = ttk.Frame(self.tab_pipe)
        outer.pack(fill="both", expand=True)

        btn_area = ttk.Frame(outer)
        btn_area.pack(fill="x", padx=40, pady=8)

        # One-click Quick Run
        self.btn_quick = tk.Button(
            btn_area, text="QUICK RUN ALL  (Extract -> Merge -> Overview)",
            bg="#1a73e8", fg="white", font=("Segoe UI", 11, "bold"),
            relief="flat", padx=10, pady=10, cursor="hand2",
            command=self.run_quick_all
        )
        self.btn_quick.pack(fill="x", pady=(0, 8))

        # Progress bar (hidden by default)
        self.progress_bar = ttk.Progressbar(btn_area, mode="determinate")

        ttk.Separator(btn_area, orient="horizontal").pack(fill="x", pady=4)
        ttk.Label(btn_area, text="Step by step:", font=("Segoe UI", 9), foreground="#5f6368").pack(anchor="w")

        # helper: button with a 1-px border frame
        def bordered_btn(parent, text, bg, fg, font_args=None, **kw):
            border_clr = "#b0bec5"
            f = tk.Frame(parent, bg=border_clr, bd=0)
            opts = dict(text=text, bg=bg, fg=fg, relief="flat", bd=0,
                        padx=8, pady=7, cursor="hand2",
                        highlightthickness=0, activebackground=bg, activeforeground=fg)
            if font_args:
                opts["font"] = font_args
            opts.update(kw)
            b = tk.Button(f, **opts)
            b.pack(fill="both", expand=True, padx=1, pady=1)
            return f, b

        # Steps 1 & 2
        step12 = ttk.Frame(btn_area)
        step12.pack(fill="x", pady=3)
        f1, self.btn1 = bordered_btn(step12, "1.  EXTRACT", "#e3f2fd", "#0d47a1",
                                     font_args=("Segoe UI", 9, "bold"), command=self.run_step1)
        f1.pack(side="left", fill="x", expand=True, padx=(0, 3))
        f2, self.btn2 = bordered_btn(step12, "2.  MERGE", "#e3f2fd", "#0d47a1",
                                     font_args=("Segoe UI", 9, "bold"), state="disabled", command=self.run_step2)
        f2.pack(side="right", fill="x", expand=True, padx=(3, 0))

        # Step 3
        f3, self.btn3 = bordered_btn(btn_area, "3.  GENERATE OVERVIEW  (Merged)",
                                     "#e3f2fd", "#0d47a1", font_args=("Segoe UI", 9, "bold"),
                                     state="disabled", command=self.run_step3)
        f3.pack(fill="x", pady=3)

        # Overview sub-buttons (compact, lighter border)
        ov_row = tk.Frame(btn_area, bg="#b0bec5")
        ov_row.pack(fill="x", pady=1)
        for txt, sfx in [("Show All NCUs", "ALL"), ("NCU 1", "NCU1"), ("NCU 2", "NCU2"), ("NCU 3", "NCU3")]:
            tk.Button(ov_row, text=txt, bg="#eceff1", fg="#37474f", relief="flat",
                      padx=6, pady=4, cursor="hand2", bd=0, highlightthickness=0,
                      font=("Segoe UI", 8),
                      command=lambda s=sfx: self.show_overview(s)).pack(side="left", fill="x", expand=True, padx=1, pady=1)

        # Steps 4 & 5
        step45 = tk.Frame(btn_area, bg="#b0bec5")
        step45.pack(fill="x", pady=3)
        f4, self.btn4 = bordered_btn(step45,
                                     "4.  GENERATE INDIVIDUAL PLOTS\n(saves PNG for every tracker)",
                                     "#e8f5e9", "#1b5e20", font_args=("Segoe UI", 9, "bold"),
                                     state="disabled", command=self.run_step4)
        f4.pack(side="left", fill="x", expand=True, padx=(0, 1))
        f5, self.btn5 = bordered_btn(step45,
                                     "5.  RUN HEALTH CHECK\n(solar-aware, priority sort)",
                                     "#ede7f6", "#4527a0", font_args=("Segoe UI", 9, "bold"),
                                     state="disabled", command=self.run_step5)
        f5.pack(side="right", fill="x", expand=True, padx=(1, 0))

        ttk.Separator(btn_area, orient="horizontal").pack(fill="x", pady=6)

        # Advanced Options (right-aligned, small)
        tk.Button(btn_area, text="Advanced Options...", bg="#f5f5f5", fg="#616161",
                  relief="solid", bd=1, padx=8, pady=4, cursor="hand2",
                  font=("Segoe UI", 8), command=self.open_advanced).pack(anchor="e", pady=(0, 6))

        # Export buttons - clearly grouped with a label
        ttk.Label(btn_area, text="Export / Reports:", font=("Segoe UI", 8), foreground="#5f6368").pack(anchor="w")

        f6, self.btn6 = bordered_btn(btn_area,
                                     "6.  EXPORT FULL PDF  (Overview + All 370 Trackers - uses saved PNGs)",
                                     "#ffebee", "#c62828", font_args=("Segoe UI", 9, "bold"),
                                     state="disabled", command=self.run_pdf_full)
        f6.pack(fill="x", pady=2)

        f7, self.btn7 = bordered_btn(btn_area,
                                     "7.  EXPORT SAMPLE PDF  (5 Random TCUs per NCU)",
                                     "#e8f5e9", "#1b5e20", font_args=("Segoe UI", 9, "bold"),
                                     state="disabled", command=self.run_pdf_random)
        f7.pack(fill="x", pady=2)

        row8 = tk.Frame(btn_area, bg="#f5f5f5")
        row8.pack(fill="x", pady=2)

        f8, self.btn8 = bordered_btn(row8,
                                     "8a.  FULL REPORT  (.docx)",
                                     "#e8eaf6", "#283593", font_args=("Segoe UI", 9, "bold"),
                                     state="disabled", command=self.run_official_report)
        f8.pack(side="left", fill="both", expand=True, padx=(0, 2))

        f8b, self.btn8b = bordered_btn(row8,
                                       "8b.  SAMPLE REPORT  (.docx - 5 TCU per NCU, ~19 pages)",
                                       "#ede7f6", "#4527a0", font_args=("Segoe UI", 9, "bold"),
                                       state="disabled", command=self.run_official_report_sample)
        f8b.pack(side="left", fill="both", expand=True, padx=(2, 0))

        # Log
        self.log_text = tk.Text(outer, height=9, bg="#202124", fg="#e8eaed",
                                font=("Consolas", 9), state="disabled", wrap="word")
        self.log_text.pack(fill="both", expand=True, padx=10, pady=(5, 8))

    def _build_dashboard_tab(self):
        # Summary banner
        banner = tk.Frame(self.tab_dash, bg="#1a73e8")
        banner.pack(fill="x")
        self.dash_label = tk.Label(banner, text="Run Health Check (Step 5) to populate this dashboard.",
                                   bg="#1a73e8", fg="white", font=("Segoe UI", 9), pady=5, padx=10,
                                   anchor="w")
        self.dash_label.pack(side="left", fill="x", expand=True)
        tk.Button(banner, text="Export CSV", bg="#1557b0", fg="white", relief="flat",
                  padx=10, pady=4, cursor="hand2", font=("Segoe UI", 8),
                  command=self.export_health_csv).pack(side="right", padx=8, pady=4)

        # Legend row
        legend = tk.Frame(self.tab_dash, bg="#f5f5f5", pady=3)
        legend.pack(fill="x")
        for label, bg, fg in [
            ("  STUCK  ", "#b71c1c", "white"),
            ("  HIGH WIND  ", "#e65100", "white"),
            ("  NO DATA  ", "#880e4f", "white"),
            ("  COMM ERROR  ", "#ef9a9a", "#333"),
            ("  SUDDEN JUMP  ", "#fff59d", "#333"),
            ("  LOW ANGLE  ", "#f1f8e9", "#333"),
        ]:
            tk.Label(legend, text=label, bg=bg, fg=fg,
                     font=("Segoe UI", 8), relief="flat", bd=0, padx=2).pack(side="left", padx=2, pady=2)
        tk.Label(legend, text="  Double-click any row to view tracker plot",
                 bg="#f5f5f5", fg="#757575", font=("Segoe UI", 8, "italic")).pack(side="right", padx=8)

        # Treeview
        cols = ("NCU", "TCU", "Type", "Severity", "Details")
        frame = tk.Frame(self.tab_dash, bg="#dadce0")
        frame.pack(fill="both", expand=True, padx=8, pady=(4, 8))

        style = ttk.Style()
        style.configure("Dashboard.Treeview",
                         rowheight=24, font=("Segoe UI", 9), background="white")
        style.configure("Dashboard.Treeview.Heading",
                         font=("Segoe UI", 9, "bold"), background="#e8eaf6", foreground="#1a237e")
        style.map("Dashboard.Treeview", background=[("selected", "#c5cae9")])

        self.tree = ttk.Treeview(frame, columns=cols, show="headings", style="Dashboard.Treeview")
        for col, width, anchor in [
            ("NCU", 52, "center"), ("TCU", 58, "center"),
            ("Type", 155, "w"), ("Severity", 145, "w"), ("Details", 0, "w")
        ]:
            self.tree.heading(col, text=col)
            self.tree.column(col, width=width, anchor=anchor, minwidth=width,
                             stretch=(col == "Details"))
        vsb = ttk.Scrollbar(frame, orient="vertical", command=self.tree.yview)
        hsb = ttk.Scrollbar(frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscroll=vsb.set, xscroll=hsb.set)
        self.tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")
        frame.rowconfigure(0, weight=1)
        frame.columnconfigure(0, weight=1)
        self.tree.bind("<Double-1>", self.on_dash_double_click)

        for tag, bg, fg in [
            ("stuck",       "#b71c1c", "white"),   # STUCK: dark red - highest priority
            ("stuck_multi", "#c62828", "white"),
            ("wind",        "#e65100", "white"),   # HIGH WIND: dark orange
            ("wind_multi",  "#bf360c", "white"),
            ("nodata",      "#880e4f", "white"),   # NO DATA: dark purple
            ("nodata_multi","#6a1b9a", "white"),
            ("comm",        "#ef9a9a", "#333"),    # COMM ERROR: light red
            ("comm_multi",  "#e57373", "#333"),
            ("jump",        "#fff59d", "#333"),    # SUDDEN JUMP: yellow
            ("jump_multi",  "#fff176", "#333"),
            ("low",         "#f1f8e9", "#333"),    # LOW ANGLE: light green
            ("low_multi",   "#dcedc8", "#333"),
        ]:
            self.tree.tag_configure(tag, background=bg, foreground=fg)

    # ---- Calendar popup ----
    def open_calendar(self):
        try:
            y = int(self.var_year.get())
            m = int(self.var_month.get())
            d = int(self.var_day.get())
        except ValueError:
            y, m, d = datetime.now().year, datetime.now().month, datetime.now().day

        def on_pick(dt):
            self.var_year.set(str(dt.year))
            self.var_month.set(str(dt.month).zfill(2))
            self.var_day.set(str(dt.day).zfill(2))

        CalendarPopup(self.root, y, m, d, on_pick)

    def open_advanced(self):
        AdvancedSettingsDialog(self.root, self)

    def browse(self):
        d = filedialog.askdirectory()
        if d:
            self.folder_path.set(d)

    def run_thread(self, target):
        threading.Thread(target=target, daemon=True).start()

    # ---- Status check ----
    def check_status(self):
        base = self.folder_path.get()
        date = self.date_val.get()
        if not base or not date or len(date) < 10:
            return
        y, m = date[:4], date[5:7]
        raw_dir = os.path.join(base, "01_Original_files", y, m, date)
        down_dir = os.path.join(base, "02_DownSampled_Files", y, m, date)
        merged = os.path.join(base, "03_Merged_files", y, m, f"{date}_1min_merged.csv")
        plots_dir = os.path.join(base, "04_Tracker_plots_angles", y, m, date, "each_tracker_plots")

        has_raw = os.path.isdir(raw_dir)
        has_down = os.path.isdir(down_dir) and bool(glob.glob(os.path.join(down_dir, "*.csv")))
        has_merged = os.path.isfile(merged)
        has_plots = os.path.isdir(plots_dir) and bool(glob.glob(os.path.join(plots_dir, "*.png")))

        def _state(condition, done=False):
            if done:
                return ("normal", "#c8e6c9", "#2e7d32")
            return ("normal", "#e8f0fe", "#1a73e8") if condition else ("disabled", "#eceff1", "#90a4ae")

        s, bg, fg = _state(has_raw)
        self.btn1.config(state=s, bg="#c8e6c9" if has_down else bg, fg="#2e7d32" if has_down else fg)
        s, bg, fg = _state(has_down)
        self.btn2.config(state=s, bg="#c8e6c9" if has_merged else bg, fg="#2e7d32" if has_merged else fg)
        s, bg, fg = _state(has_merged)
        self.btn3.config(state=s)
        self.btn4.config(state=s)
        self.btn5.config(state=s)
        self.btn6.config(state="normal" if has_plots else ("normal" if has_merged else "disabled"))
        self.btn7.config(state=s)
        self.btn8.config(state="normal" if has_merged else "disabled")
        self.btn8b.config(state="normal" if has_merged else "disabled")
        self.btn_quick.config(state="normal" if has_raw else "disabled")

    # ---- Progress bar helpers ----
    def _show_progress(self, maximum):
        self.progress_bar.config(maximum=maximum, value=0)
        self.progress_bar.pack(fill="x", pady=2)

    def _update_progress(self, value):
        self.progress_bar.config(value=value)

    def _hide_progress(self):
        self.progress_bar.pack_forget()

    # ---- Step 1: Extract ----
    def run_step1(self):
        self.run_thread(self.exec_step1)

    def exec_step1(self):
        base, date = self.folder_path.get(), self.date_val.get()
        y, m = date[:4], date[5:7]
        in_d = os.path.join(base, "01_Original_files", y, m, date)
        out_d = os.path.join(base, "02_DownSampled_Files", y, m, date)
        os.makedirs(out_d, exist_ok=True)
        print(f"\n[{datetime.now().strftime('%H:%M:%S')}] STEP 1: EXTRACT")

        # Rename .txt to .csv
        for txt in glob.glob(os.path.join(in_d, "*.txt")):
            new = os.path.splitext(txt)[0] + ".csv"
            try:
                os.rename(txt, new)
                print(f"  Renamed: {os.path.basename(txt)} -> {os.path.basename(new)}")
            except Exception as e:
                print(f"  Rename failed: {e}")

        files = glob.glob(os.path.join(in_d, "*.csv"))
        total = len(files)
        if total == 0:
            print("  No files found in source folder.")
            return

        print(f"  Found {total} files to extract.")
        self.root.after(0, lambda: self._show_progress(total))
        args_list = [(f, os.path.join(out_d, os.path.basename(f))) for f in files]
        completed = 0

        with ProcessPoolExecutor(self.cpu_cores) as ex:
            futures = {ex.submit(worker_extract, a): a[0] for a in args_list}
            for fut in as_completed(futures):
                completed += 1
                fname = os.path.basename(futures[fut])
                res = fut.result()
                parts = res.split(":")
                status = parts[0]
                rows = parts[2] if len(parts) > 2 else ""
                print(f"  [{completed:>3}/{total}] {status} - {fname}" + (f" ({rows} rows)" if rows else ""))
                c = completed
                self.root.after(0, lambda v=c: self._update_progress(v))

        self.root.after(0, self._hide_progress)
        print(f"  Extract complete ({completed}/{total} files).")
        self.root.after(0, self.check_status)

    # ---- Step 2: Merge ----
    def run_step2(self):
        self.run_thread(self.exec_step2)

    def exec_step2(self):
        base, date = self.folder_path.get(), self.date_val.get()
        y, m = date[:4], date[5:7]
        in_d = os.path.join(base, "02_DownSampled_Files", y, m, date)
        out_d = os.path.join(base, "03_Merged_files", y, m)
        os.makedirs(out_d, exist_ok=True)
        print(f"\n[{datetime.now().strftime('%H:%M:%S')}] STEP 2: MERGE")

        files = glob.glob(os.path.join(in_d, "*.csv"))
        total = len(files)
        print(f"  Found {total} files. Batch merging...")

        BATCH = 6
        chunks = []
        batches = [files[i:i + BATCH] for i in range(0, total, BATCH)]

        for i, batch in enumerate(batches):
            print(f"  Batch {i + 1}/{len(batches)} ({len(batch)} files)...")
            with ProcessPoolExecutor(self.cpu_cores) as ex:
                dfs = list(ex.map(worker_read_csv, batch))
            if not dfs:
                continue
            bd = pd.concat(dfs, ignore_index=True, copy=False)
            bd["Timestamp"] = pd.to_datetime(bd["Timestamp"], format="%d/%m/%Y %H:%M:%S.%f")
            mini = bd.set_index("Timestamp").groupby(["NCU", "TCU", "TC", "Parametro"]).resample("1min")["Valore"].mean().reset_index()
            chunks.append(mini)
            del bd, dfs
            gc.collect()

        print("  Concatenating all batches...")
        full = pd.concat(chunks, ignore_index=True, copy=False)
        final = full.groupby(["NCU", "TCU", "TC", "Parametro", "Timestamp"])["Valore"].mean().reset_index()
        out_file = os.path.join(out_d, f"{date}_1min_merged.csv")
        final.to_csv(out_file, index=False)
        print(f"  Merged CSV saved: {out_file}")
        del final, chunks, full
        gc.collect()
        self.root.after(0, self.check_status)

    # ---- Step 3: Overview ----
    def run_step3(self):
        self.run_thread(self.exec_step3)

    def exec_step3(self):
        base, date = self.folder_path.get(), self.date_val.get()
        y, m = date[:4], date[5:7]
        print(f"\n[{datetime.now().strftime('%H:%M:%S')}] STEP 3: GENERATE OVERVIEW")
        csv_p = os.path.join(base, "03_Merged_files", y, m, f"{date}_1min_merged.csv")
        out_d = os.path.join(base, "04_Tracker_plots_angles", y, m, date)
        msg = worker_overview_analysis(csv_p, out_d, date, self.angle_threshold)
        print(f"  {msg}")
        self.root.after(0, self.check_status)

    def show_overview(self, suffix):
        base, date = self.folder_path.get(), self.date_val.get()
        y, m = date[:4], date[5:7]
        path = os.path.normpath(os.path.join(base, "04_Tracker_plots_angles", y, m, date, f"NCU_TCU_{date}_{suffix}.png"))
        if not os.path.exists(path):
            print(f"Overview image not found: {path}\nRun Step 3 first.")
            return
        try:
            os.startfile(path)
        except Exception:
            import subprocess
            subprocess.run(["explorer", path], shell=True)

    # ---- Step 4: Individual plots ----
    def run_step4(self):
        self.run_thread(self.exec_step4)

    def exec_step4(self):
        base, date = self.folder_path.get(), self.date_val.get()
        y, m = date[:4], date[5:7]
        print(f"\n[{datetime.now().strftime('%H:%M:%S')}] STEP 4: INDIVIDUAL PLOTS")
        csv_p = os.path.join(base, "03_Merged_files", y, m, f"{date}_1min_merged.csv")
        out_d = os.path.join(base, "04_Tracker_plots_angles", y, m, date, "each_tracker_plots")
        os.makedirs(out_d, exist_ok=True)
        df = pd.read_csv(csv_p, dtype={"NCU": int, "TCU": int, "TC": int, "Valore": float})
        groups = [(n, t, g.to_dict("list"), os.path.join(out_d, f"TX_{n}_TCU_{int(t)}.png"))
                  for (n, t), g in df.groupby(["NCU", "TCU"])]
        total = len(groups)
        print(f"  Queuing {total} plots...")
        del df
        gc.collect()
        self.root.after(0, lambda: self._show_progress(total))
        done = [0]
        with ProcessPoolExecutor(4) as ex:
            for res in ex.map(worker_plot_file, groups):
                done[0] += 1
                d = done[0]
                if d % 20 == 0:
                    print(f"  Plots done: {d}/{total}")
                self.root.after(0, lambda v=d: self._update_progress(v))
        self.root.after(0, self._hide_progress)
        print(f"  Individual plots saved ({done[0]} files).")
        self.root.after(0, self.check_status)

    # ---- Step 5: Health check ----
    def run_step5(self):
        self.run_thread(self.exec_step5)

    def exec_step5(self):
        base, date = self.folder_path.get(), self.date_val.get()
        y, m = date[:4], date[5:7]
        print(f"\n[{datetime.now().strftime('%H:%M:%S')}] STEP 5: HEALTH CHECK")
        csv_p = os.path.join(base, "03_Merged_files", y, m, f"{date}_1min_merged.csv")
        issues = worker_health_check(csv_p, date, self.angle_threshold, self.dev_threshold)
        print(f"  Health check found {len(issues)} issues.")
        self.root.after(0, lambda: self.show_health_results(issues))

    def show_health_results(self, issues):
        self.tabs.select(self.tab_dash)
        for r in self.tree.get_children():
            self.tree.delete(r)

        # STUCK and blocked trackers have highest priority
        type_priority = {
            "STUCK": 0, "HIGH WIND MODE": 1, "NO DATA": 2,
            "COMM ERROR": 3, "SUDDEN JUMP": 4, "LOW ANGLE": 5
        }
        issues.sort(key=lambda x: (type_priority.get(x["Type"], 9), x["NCU"], x["TCU"]))

        # Find trackers with multiple issues
        counts = Counter((i["NCU"], i["TCU"]) for i in issues)
        multi = {k for k, v in counts.items() if v > 1}

        self.health_issues = {}
        type_tag_map = {
            "STUCK": "stuck", "HIGH WIND MODE": "wind",
            "NO DATA": "nodata", "COMM ERROR": "comm",
            "SUDDEN JUMP": "jump", "LOW ANGLE": "low",
        }
        for issue in issues:
            key = (issue["NCU"], issue["TCU"])
            self.health_issues.setdefault(key, []).append(issue)
            is_multi = key in multi
            tag_base = type_tag_map.get(issue["Type"], "low")
            tag = f"{tag_base}_multi" if is_multi else tag_base
            msg = issue["Msg"] + ("  [+combined]" if is_multi else "")
            self.tree.insert("", "end",
                             values=(issue["NCU"], issue["TCU"], issue["Type"], issue["Sev"], msg),
                             tags=(tag,))

        stuck_count = sum(1 for i in issues if i["Type"] == "STUCK")
        wind_count = sum(1 for i in issues if i["Type"] == "HIGH WIND MODE")
        summary = f"{len(issues)} issues"
        if stuck_count:
            summary += f"  |  {stuck_count} STUCK"
        if wind_count:
            summary += f"  |  {wind_count} HIGH WIND"
        if multi:
            summary += f"  |  {len(multi)} trackers with multiple issues [+combined]"
        self.dash_label.config(text=summary)

    # ---- Dashboard: double-click ----
    def on_dash_double_click(self, event):
        item = self.tree.selection()
        if not item:
            return
        vals = self.tree.item(item[0], "values")
        ncu, tcu = int(vals[0]), int(vals[1])
        issue_type = vals[2]
        base, date = self.folder_path.get(), self.date_val.get()
        y, m = date[:4], date[5:7]

        img_path = os.path.normpath(os.path.join(
            base, "04_Tracker_plots_angles", y, m, date,
            "each_tracker_plots", f"TX_{ncu}_TCU_{int(tcu)}.png"
        ))

        if issue_type in ("STUCK", "HIGH WIND MODE"):
            self.show_stuck_detail(ncu, tcu, img_path, date)
        elif os.path.exists(img_path):
            try:
                os.startfile(img_path)
            except Exception:
                self.show_live_plot(ncu, tcu, date)
        else:
            self.show_live_plot(ncu, tcu, date)

    def show_stuck_detail(self, ncu, tcu, img_path, date):
        """Popup with embedded plot + stuck tracker stats."""
        key = (ncu, tcu)
        issues_for_tracker = self.health_issues.get(key, [])
        stuck_issues = [i for i in issues_for_tracker if i["Type"] in ("STUCK", "HIGH WIND MODE")]

        top = tk.Toplevel(self.root)
        top.title(f"Tracker Detail: NCU {ncu} - TCU {tcu}")
        top.geometry("950x580")

        # Left: plot
        plot_frame = ttk.Frame(top)
        plot_frame.pack(side="left", fill="both", expand=True)

        base = self.folder_path.get()
        y, m = date[:4], date[5:7]
        csv_p = os.path.join(base, "03_Merged_files", y, m, f"{date}_1min_merged.csv")

        try:
            df = pd.read_csv(csv_p, dtype={"NCU": int, "TCU": int, "TC": int, "Valore": float})
            df["Timestamp"] = pd.to_datetime(df["Timestamp"])
            sub = df[(df["NCU"] == ncu) & (df["TCU"] == tcu)]
            target = sub[sub["TC"] == 4].sort_values("Timestamp")
            actual = sub[sub["TC"] == 5].sort_values("Timestamp")

            fig = Figure(figsize=(7, 4.5), dpi=100)
            ax = fig.add_subplot(111)
            if not target.empty:
                ax.plot(target["Timestamp"], target["Valore"], label="Target", color="royalblue", lw=1)
            if not actual.empty:
                ax.plot(actual["Timestamp"], actual["Valore"], label="Actual", color="tomato", lw=1.2)

            # Highlight stuck periods
            for si in stuck_issues:
                if "StuckStart" in si and "StuckEnd" in si:
                    ax.axvspan(si["StuckStart"], si["StuckEnd"], alpha=0.15, color="red")

            ax.set_title(f"NCU {ncu} - TCU {tcu}")
            ax.set_ylabel("Angle (deg)")
            ax.legend(fontsize=8)
            ax.grid(True, alpha=0.3)
            fig.tight_layout()

            canvas = FigureCanvasTkAgg(fig, master=plot_frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill="both", expand=True, padx=5, pady=5)
            del df, sub
        except Exception as e:
            ttk.Label(plot_frame, text=f"Could not load plot:\n{e}").pack(padx=20, pady=20)

        # Right: stats panel
        info_frame = tk.Frame(top, bg="#f8f9fa", width=240)
        info_frame.pack(side="right", fill="y", padx=5, pady=5)
        info_frame.pack_propagate(False)

        tk.Label(info_frame, text=f"NCU {ncu}  |  TCU {tcu}",
                 bg="#1a73e8", fg="white", font=("Segoe UI", 11, "bold"),
                 pady=8).pack(fill="x")

        for issue in issues_for_tracker:
            color = {"Communication Error": "#ffcdd2", "Critical": "#ffcdd2",
                     "High": "#ffe0b2", "Medium": "#fff9c4", "Low": "#f1f8e9"}.get(issue["Sev"], "white")
            block = tk.Frame(info_frame, bg=color, padx=8, pady=6)
            block.pack(fill="x", padx=5, pady=4)
            tk.Label(block, text=issue["Type"], bg=color, font=("Segoe UI", 9, "bold")).pack(anchor="w")
            tk.Label(block, text=f"Severity: {issue['Sev']}", bg=color, font=("Segoe UI", 8)).pack(anchor="w")
            tk.Label(block, text=issue["Msg"], bg=color, font=("Segoe UI", 8), wraplength=200, justify="left").pack(anchor="w")
            if "StuckAngle" in issue:
                tk.Label(block, text=f"Stuck at: {issue['StuckAngle']:.1f} deg", bg=color, font=("Segoe UI", 8)).pack(anchor="w")
            if "AvgDev" in issue:
                tk.Label(block, text=f"Avg deviation: {issue['AvgDev']:.1f} deg", bg=color, font=("Segoe UI", 8)).pack(anchor="w")
            if "StuckStart" in issue:
                t0 = issue["StuckStart"]
                t1 = issue["StuckEnd"]
                tk.Label(block, text=f"From: {t0.strftime('%H:%M')}", bg=color, font=("Segoe UI", 8)).pack(anchor="w")
                tk.Label(block, text=f"To:   {t1.strftime('%H:%M')}", bg=color, font=("Segoe UI", 8)).pack(anchor="w")

        ttk.Button(info_frame, text="Open saved PNG",
                   command=lambda: os.startfile(img_path) if os.path.exists(img_path) else None
                   ).pack(pady=10)

    def show_live_plot(self, ncu, tcu, date):
        base = self.folder_path.get()
        y, m = date[:4], date[5:7]
        csv_p = os.path.join(base, "03_Merged_files", y, m, f"{date}_1min_merged.csv")
        if not os.path.isfile(csv_p):
            messagebox.showerror("Error", f"Merged CSV not found:\n{csv_p}")
            return
        try:
            df = pd.read_csv(csv_p, dtype={"NCU": int, "TCU": int, "TC": int, "Valore": float})
            df["Timestamp"] = pd.to_datetime(df["Timestamp"])
            sub = df[(df["NCU"] == ncu) & (df["TCU"] == tcu)]
            target = sub[sub["TC"] == 4].sort_values("Timestamp")
            actual = sub[sub["TC"] == 5].sort_values("Timestamp")

            top = tk.Toplevel(self.root)
            top.title(f"Live Plot: NCU {ncu} - TCU {tcu}")
            top.geometry("900x480")
            fig = Figure(figsize=(8, 4), dpi=100)
            ax = fig.add_subplot(111)
            if not target.empty:
                ax.plot(target["Timestamp"], target["Valore"], label="Target", color="royalblue")
            if not actual.empty:
                ax.plot(actual["Timestamp"], actual["Valore"], label="Actual", color="tomato", linestyle="--")
            ax.set_title(f"NCU {ncu} - TCU {tcu}")
            ax.set_ylabel("Angle (deg)")
            ax.legend()
            ax.grid(True, alpha=0.3)
            if not actual.empty:
                stats = f"Min: {actual['Valore'].min():.1f}  Max: {actual['Valore'].max():.1f}  Avg: {actual['Valore'].mean():.1f}"
                ax.text(0.02, 0.02, stats, transform=ax.transAxes, fontsize=9,
                        bbox=dict(boxstyle="round,pad=0.3", facecolor='white', alpha=0.8))
            fig.tight_layout()
            FigureCanvasTkAgg(fig, master=top).get_tk_widget().pack(fill="both", expand=True)
            del df, sub, target, actual
        except Exception as e:
            messagebox.showerror("Plot Error", str(e))

    # ---- Export PDF buttons ----
    def run_pdf_full(self):
        self.run_thread(self.exec_pdf_full)

    def exec_pdf_full(self):
        base, date = self.folder_path.get(), self.date_val.get()
        y, m = date[:4], date[5:7]
        print(f"\n[{datetime.now().strftime('%H:%M:%S')}] STEP 6: FULL PDF EXPORT")
        overview_dir = os.path.join(base, "04_Tracker_plots_angles", y, m, date)
        plots_dir = os.path.join(overview_dir, "each_tracker_plots")
        out_pdf = os.path.join(base, "05_Tracker_Report_PDF", y, m, f"Tracker_Report_{date}.pdf")
        os.makedirs(os.path.dirname(out_pdf), exist_ok=True)

        if not os.path.isdir(plots_dir) or not glob.glob(os.path.join(plots_dir, "*.png")):
            print("  No saved PNG images found. Run Step 4 (Generate Individual Plots) first.")
            print("  Falling back to data-based PDF generation...")
            # Fallback: use merged CSV directly
            from matplotlib.backends.backend_pdf import PdfPages
            import matplotlib.image as mpimg
            csv_p = os.path.join(base, "03_Merged_files", y, m, f"{date}_1min_merged.csv")
            if not os.path.isfile(csv_p):
                self.root.after(0, lambda: messagebox.showerror("Error", "Merged CSV not found. Run steps 1-3 first."))
                return
            df = pd.read_csv(csv_p, dtype={"NCU": int, "TCU": int, "TC": int, "Valore": float})
            df["Timestamp"] = pd.to_datetime(df["Timestamp"])
            with PdfPages(out_pdf) as pdf:
                for sfx in ["ALL", "NCU1", "NCU2", "NCU3"]:
                    ip = os.path.join(overview_dir, f"NCU_TCU_{date}_{sfx}.png")
                    if os.path.exists(ip):
                        img = mpimg.imread(ip)
                        fig, ax = plt.subplots(figsize=(15.5, 11.2))
                        ax.imshow(img); ax.axis('off')
                        pdf.savefig(fig, orientation='landscape', bbox_inches='tight', dpi=150)
                        plt.close(fig); del img
                for (ncu, tcu), g in df.groupby(["NCU", "TCU"]):
                    t = g[g["TC"] == 4].sort_values("Timestamp")
                    a = g[g["TC"] == 5].sort_values("Timestamp")
                    fig, ax = plt.subplots(figsize=(10, 6))
                    if not t.empty: ax.plot(t["Timestamp"], t["Valore"], label="Target", color="royalblue", lw=1)
                    if not a.empty: ax.plot(a["Timestamp"], a["Valore"], label="Actual", color="tomato", lw=1)
                    ax.set_title(f"NCU {ncu} - TCU {tcu}"); ax.grid(True, alpha=0.3); ax.legend()
                    pdf.savefig(fig, dpi=100); plt.close(fig)
            del df
            gc.collect()
            msg = f"PDF (data-based) saved to:\n{out_pdf}"
        else:
            msg = worker_generate_pdf_from_images(overview_dir, plots_dir, out_pdf, date)

        print(f"  {msg}")
        self.root.after(0, lambda: messagebox.showinfo("PDF Export", msg))

    def run_pdf_random(self):
        self.run_thread(self.exec_pdf_random)

    def exec_pdf_random(self):
        base, date = self.folder_path.get(), self.date_val.get()
        y, m = date[:4], date[5:7]
        print(f"\n[{datetime.now().strftime('%H:%M:%S')}] STEP 7: SAMPLE PDF EXPORT")
        csv_p = os.path.join(base, "03_Merged_files", y, m, f"{date}_1min_merged.csv")
        img_folder = os.path.join(base, "04_Tracker_plots_angles", y, m, date, "each_tracker_plots")
        out_pdf = os.path.join(base, "05_Tracker_Report_PDF", y, m, f"Tracker_Sample_{date}.pdf")
        os.makedirs(os.path.dirname(out_pdf), exist_ok=True)
        msg = worker_generate_random_pdf(csv_p, img_folder, out_pdf, date)
        print(f"  {msg}")
        self.root.after(0, lambda: messagebox.showinfo("Sample PDF", msg))

    def run_official_report(self):
        self.run_thread(self.exec_official_report)

    def exec_official_report(self):
        base, date = self.folder_path.get(), self.date_val.get()
        y, m = date[:4], date[5:7]
        print(f"\n[{datetime.now().strftime('%H:%M:%S')}] STEP 8: OFFICIAL DOCX REPORT")

        template_path = os.path.join(base, TEMPLATE_FILENAME)
        if not os.path.isfile(template_path):
            msg = f"Template not found:\n{template_path}\n\nPlace the template file in the root folder."
            print(f"  ERROR: {msg}")
            self.root.after(0, lambda: messagebox.showerror("Template Missing", msg))
            return

        dt = datetime.strptime(date, "%Y-%m-%d")
        out_name = f"{dt.strftime('%Y.%m.%d')} - Report Tracker.docx"
        out_path = os.path.join(base, "05_Tracker_Report_PDF", y, m, out_name)
        os.makedirs(os.path.dirname(out_path), exist_ok=True)

        msg = worker_generate_official_report(base, date, template_path, out_path, sample_n=0)
        print(f"  {msg}")
        self.root.after(0, lambda: messagebox.showinfo("Official Report", msg))

    def run_official_report_sample(self):
        self.run_thread(self.exec_official_report_sample)

    def exec_official_report_sample(self):
        base, date = self.folder_path.get(), self.date_val.get()
        y, m = date[:4], date[5:7]
        print(f"\n[{datetime.now().strftime('%H:%M:%S')}] STEP 8b: SAMPLE DOCX REPORT (5 TCU per NCU)")

        template_path = os.path.join(base, TEMPLATE_FILENAME)
        if not os.path.isfile(template_path):
            msg = f"Template not found:\n{template_path}\n\nPlace the template file in the root folder."
            print(f"  ERROR: {msg}")
            self.root.after(0, lambda: messagebox.showerror("Template Missing", msg))
            return

        dt = datetime.strptime(date, "%Y-%m-%d")
        out_name = f"{dt.strftime('%Y.%m.%d')} - Report Tracker SAMPLE.docx"
        out_path = os.path.join(base, "05_Tracker_Report_PDF", y, m, out_name)
        os.makedirs(os.path.dirname(out_path), exist_ok=True)

        msg = worker_generate_official_report(base, date, template_path, out_path, sample_n=5)
        print(f"  {msg}")
        self.root.after(0, lambda: messagebox.showinfo("Sample Report", msg))

    # ---- Quick run all ----
    def run_quick_all(self):
        def _run():
            print(f"\n[{datetime.now().strftime('%H:%M:%S')}] === QUICK RUN: Extract -> Merge -> Overview ===")
            self.exec_step1()
            self.exec_step2()
            self.exec_step3()
            print(f"[{datetime.now().strftime('%H:%M:%S')}] === QUICK RUN COMPLETE ===")
        self.run_thread(_run)

    # ---- Health CSV export ----
    def export_health_csv(self):
        if not self.health_issues:
            messagebox.showinfo("No Data", "Run Health Check first.")
            return
        path = filedialog.asksaveasfilename(defaultextension=".csv", filetypes=[("CSV", "*.csv")])
        if not path:
            return
        rows = []
        for issues in self.health_issues.values():
            for i in issues:
                rows.append({k: v for k, v in i.items() if k in ("NCU", "TCU", "Type", "Sev", "Msg")})
        pd.DataFrame(rows).to_csv(path, index=False)
        messagebox.showinfo("Exported", f"Health report saved to:\n{path}")


# ==============================================================================
# ENTRY POINT
# ==============================================================================

if __name__ == "__main__":
    multiprocessing.freeze_support()
    root = tk.Tk()
    app = TrackerSuiteApp(root)
    root.mainloop()
